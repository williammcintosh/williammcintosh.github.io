from flask import Flask, render_template, request, send_file
import requests
import openai
from dotenv import load_dotenv
import os
import logging
from bs4 import BeautifulSoup
import re
from datetime import datetime
import tempfile  # To manage temporary files
from urllib.parse import urlparse  # For URL parsing
from docx import Document  # For Word document creation

# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.config['WTF_CSRF_ENABLED'] = False  # Disable CSRF protection

# Get OpenAI API key from environment variable
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    raise ValueError("No API key found. Please set the OPENAI_API_KEY environment variable.")

openai.api_key = api_key

# Suppress detailed stacktrace in the output
logging.getLogger('werkzeug').setLevel(logging.ERROR)

def get_completion(prompt, model="gpt-4-1106-preview"):
    completion = openai.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": prompt,
            },
        ],
    )
    return completion.choices[0].message.content

def fetch_full_page_source(url):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.121 Safari/537.36",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
    }

    # Use a session to manage cookies and headers
    session = requests.Session()
    session.headers.update(headers)

    try:
        response = session.get(url)
        response.raise_for_status()  # Raise an error for bad responses like 429

        return response.text
    except requests.exceptions.HTTPError as err:
        return f"HTTP error occurred: {err}"
    except requests.exceptions.RequestException as e:
        return f"An error occurred: {str(e)}"

def extract_main_content(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    main_content = soup.find_all(['p', 'h1', 'h2', 'h3', 'li'])
    content_text = "\n".join([element.get_text() for element in main_content if element.get_text().strip() != ""])
    return content_text

def extract_employer_name(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Find the specific div containing the employer name
    employer_name_element = soup.find('div', class_='job-details-jobs-unified-top-card__company-name')
    
    # Get the text inside the anchor tag within the div
    if employer_name_element:
        employer_name = employer_name_element.find('a')
        return employer_name.get_text().strip() if employer_name else "Unknown"
    
    return "Unknown"

def extract_hiring_manager_email(html_content):
    # Use regex to find email addresses directly in the HTML content
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, html_content)
    
    # If emails are found, return the first one
    if emails:
        return emails[0]
    return None

def generate_and_save_CV_latex(job_posting_html):
    # Read the prompt and JSON experience from text files
    with open('templates/cv_prompt.txt', 'r') as file:
        cv_prompt = file.read()

    with open('templates/experience_latex.txt', 'r') as file:
        experience_cv = file.read()

    prompt = cv_prompt.replace('◊', '\\') + job_posting_html + experience_cv.replace('◊', '\\')

    latex_code = get_completion(prompt)

    # Save LaTeX code to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".tex") as tmp_latex:
        tmp_latex.write(latex_code.encode('utf-8'))
        tmp_latex_path = tmp_latex.name

    return tmp_latex_path

def generate_and_save_Coverletter_latex(html_content):
    # Read the cover letter prompt from a text file
    with open('templates/coverletter_prompt.txt', 'r') as file:
        prompt = file.read()
    
    # Create the LaTeX cover letter using OpenAI's API
    latex_code = get_completion(prompt)

    # Save LaTeX code to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".tex") as tmp_latex:
        tmp_latex.write(latex_code.encode('utf-8'))
        tmp_latex_path = tmp_latex.name

    return tmp_latex_path

def save_as_word(html_content):
    """Converts HTML content to a Word document."""
    document = Document()
    document.add_heading('Job Information', level=1)

    # Add HTML content to the Word document
    soup = BeautifulSoup(html_content, 'html.parser')
    main_content = soup.find_all(['p', 'h1', 'h2', 'h3', 'li'])
    for element in main_content:
        text = element.get_text()
        if element.name in ['h1', 'h2', 'h3']:
            document.add_heading(text, level=int(element.name[1]))
        else:
            document.add_paragraph(text)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_doc:
        document.save(tmp_doc.name)
        return tmp_doc.name

def extract_source(url):
    """Extracts the source platform from the URL."""
    parsed_url = urlparse(url)
    domain = parsed_url.netloc

    # Check common platforms
    if "linkedin" in domain:
        return "LinkedIn"
    elif "seek" in domain:
        return "Seek"
    else:
        # Default to domain name for unknown sources
        return domain.split('.')[0].capitalize()

def submit_google_form(year, month, day, company, role, location, contact, source):
    form_url = "https://docs.google.com/forms/u/0/d/e/1FAIpQLSfCP-fBMn9go2PmCcfz2soJ7qeqv_vuHE1Fvf9x-sg-iP2kDQ/formResponse"
    form_data = {
        "entry.1298833329_year": year,  # Year part of the date
        "entry.1298833329_month": month,  # Month part of the date
        "entry.1298833329_day": day,  # Day part of the date
        "entry.1710074163": company,
        "entry.302644938": role,
        "entry.1204204715": location,
        "entry.1577944511": contact,
        "entry.350151551": source  # Source field from Google Form
    }
    response = requests.post(form_url, data=form_data)
    return response.status_code == 200

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_html', methods=['POST'])
def get_html():
    url = request.form['url']
    try:
        # Fetch full page source with requests
        html_content = fetch_full_page_source(url)

        if html_content.startswith("HTTP error occurred") or html_content.startswith("An error occurred"):
            return html_content  # Return error message if fetching fails

        # Extract main content
        extracted_content = extract_main_content(html_content)

        # Extract employer name
        employer_name = extract_employer_name(html_content)

        # Extract hiring manager email
        hiring_manager_email = extract_hiring_manager_email(html_content)

        # Extract job role from content (if necessary)
        job_role = "Data Engineer"  # Example placeholder, adjust as needed

        # Extract job location from content (if necessary)
        job_location = "Dunedin or Christchurch"  # Example placeholder, adjust as needed

        # Extract source platform
        source = extract_source(url)

        # Submit form to Google Form
        current_date = datetime.now().strftime("%Y-%m-%d")
        year, month, day = current_date.split('-')
        form_submitted = submit_google_form(year, month, day, employer_name, job_role, job_location, hiring_manager_email, source)

        # Generate and save documents
        job_description_word_path = save_as_word(html_content)
        latex_file_path = generate_and_save_CV_latex(html_content)
        coverletter_file_path = generate_and_save_Coverletter_latex(html_content)

        return render_template(
            'result.html',
            url=url,
            job_role=job_role,
            extracted_content=extracted_content,
            employer_name=employer_name,
            hiring_manager_email=hiring_manager_email or "",
            job_location=job_location,
            form_submitted=form_submitted,
            source=source,
            job_description_word_path=job_description_word_path,
            latex_file_path=latex_file_path,
            coverletter_file_path=coverletter_file_path
        )

    except Exception as e:
        return f"An error occurred: {str(e)}"

@app.route('/download/<file_type>')
def download_file(file_type):
    file_path = request.args.get('path')
    company_name = request.args.get('company', 'Unknown')
    role_name = request.args.get('role', 'Unknown')

    # Handle missing parameters
    if not file_path or not company_name or not role_name:
        return "Invalid parameters", 400

    # Clean company and role names
    company_clean = re.sub(r'[^\w]', '', company_name)[:30]
    role_clean = re.sub(r'[^\w]', '', role_name)[:30]

    # Construct the download filename based on the file type
    if file_type == "word":
        filename = f"Job_Description_{company_clean}_{role_clean}.docx"
    elif file_type == "latex":
        filename = f"McIntosh__Resume_{company_clean}_{role_clean}.tex"
    elif file_type == "coverletter":
        filename = f"McIntosh__Coverletter_{company_clean}_{role_clean}.docx"
    else:
        return "Invalid file type requested", 400

    return send_file(file_path, as_attachment=True, download_name=filename)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True)
